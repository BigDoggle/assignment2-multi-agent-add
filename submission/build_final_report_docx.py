from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Mm


ROOT = Path(__file__).resolve().parent
INPUT = ROOT / "final-report.md"
OUTPUT = ROOT / "final-report.docx"


def set_font(run, name: str, size: int, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 20), (2, 16), (3, 14)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0) if level < 3 else RGBColor(67, 67, 67)
        style.paragraph_format.space_before = Pt(20 if level == 1 else 18 if level == 2 else 16)
        style.paragraph_format.space_after = Pt(6 if level < 3 else 4)


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, "Courier New", 9)


def add_inline_runs(paragraph, text: str, default_size: int = 11) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            set_font(run, "Arial", default_size, bold=True)
            continue

        if part.startswith("`") and part.endswith("`") and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            set_font(run, "Courier New", max(default_size - 1, 9))
            continue

        run = paragraph.add_run(part)
        set_font(run, "Arial", default_size)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text, 11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text, 11)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text, 11)


def build_docx() -> None:
    if OUTPUT.exists():
        OUTPUT.unlink()
    doc = Document()
    style_document(doc)

    lines = INPUT.read_text(encoding="utf-8").splitlines()
    in_code = False

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
            else:
                in_code = False
                doc.add_paragraph()
            continue

        if in_code:
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            set_font(run, "Arial", 26, bold=False, color="000000")
            continue

        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            continue

        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        if line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            continue

        if len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            add_numbered(doc, line[3:].strip())
            continue

        add_paragraph(doc, line)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
